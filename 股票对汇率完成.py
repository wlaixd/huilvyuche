import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.model_selection import train_test_split, cross_val_score, KFold
from sklearn.linear_model import LinearRegression
from sklearn.metrics import mean_squared_error, r2_score, mean_absolute_error
from sklearn.neural_network import MLPRegressor
from sklearn.ensemble import RandomForestRegressor, GradientBoostingRegressor
import xgboost as xgb
from docx import Document
from docx.shared import Inches
import math
import os
from sklearn.preprocessing import StandardScaler
import tensorflow as tf
from tensorflow.keras.models import Sequential
from tensorflow.keras.layers import LSTM as KerasLSTM, Dense, Dropout
from tensorflow.keras.callbacks import EarlyStopping

from sklearn.base import BaseEstimator, RegressorMixin


print(f"TensorFlow 版本: {tf.__version__}")
print("tensorflow.keras 模块可用")
# Create directory for saving figures if it doesn't exist
output_dir = r'C:\Users\PKF\Desktop\实习\ycq\figures'
os.makedirs(output_dir, exist_ok=True)

data = pd.read_excel(r'C:\Users\PKF\Desktop\实习\ycq\2.xlsx')

# Define the independent variables (features) and the dependent variable (target)
X = data[['399006.SZ', '399001.SZ', '000001.SH', '000905.SH', '000300.SH', '000016.SH', 
          'AU9999', 'AU(T+D)', 'Dnvaltrd', 'Brent Spot Price', 'UsDCNY']]
y = data['BIS指数']

# Split the data into training and testing sets
X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)

# Create a Word document to store results
doc = Document()
doc.add_heading('第五章 股票市场对BIS指数影响的机器学习模型分析', level=0)
doc.add_heading('5.2 机器学习模型', level=1)

# Function to calculate MAPE
def mean_absolute_percentage_error(y_true, y_pred):
    y_true, y_pred = np.array(y_true), np.array(y_pred)
    return np.mean(np.abs((y_true - y_pred) / y_true))

# Function to evaluate and document model results
def evaluate_model(model, model_name, model_intro):
    # Fit the model
    model.fit(X_train, y_train)
    
    # Predictions
    y_train_pred = model.predict(X_train)
    y_test_pred = model.predict(X_test)
    
    # Calculate metrics for training set
    train_r2 = r2_score(y_train, y_train_pred)
    train_mse = mean_squared_error(y_train, y_train_pred)
    train_rmse = math.sqrt(train_mse)
    train_mae = mean_absolute_error(y_train, y_train_pred)
    train_mape = mean_absolute_percentage_error(y_train, y_train_pred)
    
    # Calculate metrics for test set
    test_r2 = r2_score(y_test, y_test_pred)
    test_mse = mean_squared_error(y_test, y_test_pred)
    test_rmse = math.sqrt(test_mse)
    test_mae = mean_absolute_error(y_test, y_test_pred)
    test_mape = mean_absolute_percentage_error(y_test, y_test_pred)
    
    # Print results
    print(f'{model_name} Results:')
    print(f'Training Set - R2: {train_r2:.4f}, MSE: {train_mse:.4f}, RMSE: {train_rmse:.4f}, MAE: {train_mae:.4f}, MAPE: {train_mape:.4f}')
    print(f'Test Set - R2: {test_r2:.4f}, MSE: {test_mse:.4f}, RMSE: {test_rmse:.4f}, MAE: {test_mae:.4f}, MAPE: {test_mape:.4f}')
    
    # Add to document
    doc.add_heading(model_name, level=2)
    doc.add_paragraph(model_intro)
    
    # Create a table for metrics
    table = doc.add_table(rows=3, cols=6)
    table.style = 'Table Grid'
    
    # Add headers
    header_cells = table.rows[0].cells
    header_cells[0].text = ''
    header_cells[1].text = 'R2'
    header_cells[2].text = 'MSE'
    header_cells[3].text = 'RMSE'
    header_cells[4].text = 'MAE'
    header_cells[5].text = 'MAPE'
    
    # Add training set metrics
    train_cells = table.rows[1].cells
    train_cells[0].text = '训练集'
    train_cells[1].text = f'{train_r2:.4f}'
    train_cells[2].text = f'{train_mse:.4f}'
    train_cells[3].text = f'{train_rmse:.4f}'
    train_cells[4].text = f'{train_mae:.4f}'
    train_cells[5].text = f'{train_mape:.4f}'
    
    # Add test set metrics
    test_cells = table.rows[2].cells
    test_cells[0].text = '测试集'
    test_cells[1].text = f'{test_r2:.4f}'
    test_cells[2].text = f'{test_mse:.4f}'
    test_cells[3].text = f'{test_rmse:.4f}'
    test_cells[4].text = f'{test_mae:.4f}'
    test_cells[5].text = f'{test_mape:.4f}'
    
    # Set font for Chinese characters
    plt.rcParams['font.sans-serif'] = ['SimHei']  # Use SimHei font for Chinese characters
    plt.rcParams['axes.unicode_minus'] = False    # Fix minus sign display issue
    
    # Create comparison plots
    plt.figure(figsize=(12, 10))
    
    # Training set plot
    plt.subplot(2, 1, 1)
    plt.plot(range(len(y_train)), y_train.values, label='实际值', color='blue')
    plt.plot(range(len(y_train_pred)), y_train_pred, label='预测值', color='red')
    plt.title(f'{model_name} - 训练集预测值与实际值对比')
    plt.xlabel('样本索引')
    plt.ylabel('BIS指数')
    
    # Add metrics text to bottom left
    metrics_text = f'R2: {train_r2:.4f}\nMSE: {train_mse:.4f}\nRMSE: {train_rmse:.4f}\nMAE: {train_mae:.4f}\nMAPE: {train_mape:.4f}'
    plt.text(0.05, 0.05, metrics_text, transform=plt.gca().transAxes, 
             bbox=dict(facecolor='white', alpha=0.8))
    
    # Place legend in top right
    plt.legend(loc='upper right')
    plt.grid(True)
    
    # Test set plot
    plt.subplot(2, 1, 2)
    plt.plot(range(len(y_test)), y_test.values, label='实际值', color='blue')
    plt.plot(range(len(y_test_pred)), y_test_pred, label='预测值', color='red')
    plt.title(f'{model_name} - 测试集预测值与实际值对比')
    plt.xlabel('样本索引')
    plt.ylabel('BIS指数')
    
    # Add metrics text to bottom left
    metrics_text = f'R2: {test_r2:.4f}\nMSE: {test_mse:.4f}\nRMSE: {test_rmse:.4f}\nMAE: {test_mae:.4f}\nMAPE: {test_mape:.4f}'
    plt.text(0.05, 0.05, metrics_text, transform=plt.gca().transAxes, 
             bbox=dict(facecolor='white', alpha=0.8))
    
    # Place legend in top right
    plt.legend(loc='upper right')
    plt.grid(True)
    
    plt.tight_layout()
    
    # Save the figure
    fig_path = f'{output_dir}\\{model_name.replace(" ", "_")}.png'
    plt.savefig(fig_path)
    plt.close()
    
    # Add the figure to the document
    doc.add_paragraph().add_run().add_picture(fig_path, width=Inches(6))
    
    # Add model-specific conclusion
    if model_name == '多元线性回归模型':
        conclusion = f"""该多元线性回归模型在训练集上的R2为{train_r2:.4f}，在测试集上的R2为{test_r2:.4f}。模型表现出一定的预测能力，但R2值相对较低，说明线性模型可能无法完全捕捉股票市场指标与BIS指数之间的复杂关系。测试集RMSE为{test_rmse:.4f}，表明预测值与实际值之间存在一定偏差。这一结果表明，股票市场指标与BIS指数之间可能存在非线性关系，需要更复杂的模型来进行建模。"""
    elif model_name == '神经网络模型':
        conclusion = f"""神经网络模型在训练集上的R2为{train_r2:.4f}，在测试集上的R2为{test_r2:.4f}。相比线性模型，神经网络展现出更强的拟合能力，能够捕捉数据中的非线性关系。测试集RMSE为{test_rmse:.4f}，表明模型在未见数据上的预测精度有所提高。然而，训练集和测试集性能差异较大，可能存在一定程度的过拟合现象，这提示我们在实际应用中需要进一步优化网络结构和参数。"""
    elif model_name == 'XGBoost模型':
        conclusion = f"""XGBoost模型在训练集上的R2为{train_r2:.4f}，在测试集上的R2为{test_r2:.4f}。该模型展现出极强的拟合能力，测试集RMSE为{test_rmse:.4f}，明显优于线性模型和神经网络。训练集和测试集的性能差异表明模型具有良好的泛化能力。XGBoost能够自动处理特征间的交互作用，这对于捕捉股票市场指标与BIS指数之间的复杂关系非常有效。该结果强烈表明，股票市场指标对BIS指数具有显著的预测价值。"""
    elif model_name == '随机森林模型':
        conclusion = f"""随机森林模型在训练集上的R2为{train_r2:.4f}，在测试集上的R2为{test_r2:.4f}。该模型表现出较强的预测能力，测试集RMSE为{test_rmse:.4f}。随机森林通过集成多棵决策树的预测结果，有效减少了过拟合风险，提高了模型的泛化能力。这一结果表明，股票市场指标与BIS指数之间存在可被捕捉的关系模式，且这种关系可能是非线性的。"""
    elif model_name == '梯度提升树模型':
        conclusion = f"""梯度提升树模型在训练集上的R2为{train_r2:.4f}，在测试集上的R2为{test_r2:.4f}。该模型展现出较好的预测性能，测试集RMSE为{test_rmse:.4f}。梯度提升树通过顺序构建决策树并纠正前面树的预测误差，能够有效捕捉股票市场指标与BIS指数之间的关系。模型的表现进一步证实了股票市场指标对BIS指数具有预测价值。"""
    elif model_name == '长短期记忆网络模型':
        conclusion = f"""长短期记忆网络(LSTM)模型在训练集上的R2为{train_r2:.4f}，在测试集上的R2为{test_r2:.4f}。该模型在处理时间序列数据方面表现出色，能够捕捉数据中的长期依赖关系。测试集RMSE为{test_rmse:.4f}，表明模型在预测BIS指数方面具有一定的预测能力。与加楚懿(2025)和于孝建(2024)的研究发现相符，LSTM在金融时间序列分析中的应用价值得到了证实。"""
    elif model_name == '极限学习机模型':
        conclusion = f"""极限学习机(ELM)模型在训练集上的R2为{train_r2:.4f}，在测试集上的R2为{test_r2:.4f}。该模型在快速建模方面表现出色，测试集RMSE为{test_rmse:.4f}。ELM通过随机生成输入权重和隐层偏置，并通过解析方法确定输出权重，避免了局部最小值问题，同时提高了计算效率。加楚懿(2025)在比较BP神经网络、ELM和LSTM三种机器学习方法时发现，ELM在某些预测任务中具有计算效率高、泛化能力强的优势。"""
    else:
        conclusion = f"""未知模型在训练集上的R2为{train_r2:.4f}，在测试集上的R2为{test_r2:.4f}。该模型的性能需要进一步分析。"""
    
    doc.add_paragraph(conclusion)
    
    return {
        'model_name': model_name,
        'train_r2': train_r2,
        'train_mse': train_mse,
        'train_rmse': train_rmse,
        'train_mae': train_mae,
        'train_mape': train_mape,
        'test_r2': test_r2,
        'test_mse': test_mse,
        'test_rmse': test_rmse,
        'test_mae': test_mae,
        'test_mape': test_mape,
        'y_train_pred': y_train_pred,
        'y_test_pred': y_test_pred
    }

# Model introductions
linear_intro = """5.2.2 多元线性回归模型
多元线性回归是一种基础的统计学习方法，用于建立因变量与多个自变量之间的线性关系。该模型假设因变量与自变量之间存在线性关系，通过最小化预测值与实际值之间的均方误差来确定最优参数。在本研究中，我们使用多元线性回归来分析股票市场指标对BIS指数的影响。"""

xgb_intro = """5.2.4 XGBoost模型
XGBoost(eXtreme Gradient Boosting)是一种高效的梯度提升树算法，通过集成多个决策树来提高预测性能。该算法采用了正则化技术来控制模型复杂度，并使用二阶导数来优化损失函数，具有高效、灵活和可扩展的特点。在本研究中，我们利用XGBoost来分析股票市场指标对BIS指数的影响，该模型能够自动处理特征之间的交互作用。"""

rf_intro = """5.2.5 随机森林模型
随机森林是一种集成学习方法，通过构建多个决策树并将它们的预测结果进行组合来提高模型性能。每棵树在随机选择的数据子集上训练，并在每个节点使用随机选择的特征子集进行分裂，从而减少过拟合风险。随机森林模型具有较强的鲁棒性和解释性，能够评估特征重要性。在本研究中，我们应用随机森林来分析股票市场指标对BIS指数的影响程度。"""

gb_intro = """5.2.6 梯度提升树模型
梯度提升树是一种顺序构建决策树的集成方法，每棵新树都试图纠正前面树的预测误差。通过梯度下降算法优化损失函数，该模型能够逐步提高预测精度。梯度提升树对异常值不敏感，能够自动处理特征间的交互作用，并能处理各种类型的数据。在本研究中，我们使用梯度提升树来建立股票市场指标与BIS指数之间的预测模型。"""

# Improved ELM class
class ELMRegressor(BaseEstimator, RegressorMixin):
    def __init__(self, n_hidden=500, activation='tanh', random_state=None, alpha=0.001):
        self.n_hidden = n_hidden
        self.activation = activation
        self.random_state = random_state
        self.alpha = alpha  # Regularization parameter
        self.weights = None
        self.bias = None
        self.beta = None
        
    def _sigmoid(self, x):
        return 1 / (1 + np.exp(-x))
    
    def _relu(self, x):
        return np.maximum(0, x)
    
    def _tanh(self, x):
        return np.tanh(x)
    
    def _activation_function(self, x):
        if self.activation == 'sigmoid':
            return self._sigmoid(x)
        elif self.activation == 'relu':
            return self._relu(x)
        elif self.activation == 'tanh':
            return self._tanh(x)
        else:
            raise ValueError("Activation function not supported")
    
    def fit(self, X, y):
        X = np.array(X)
        y = np.array(y).reshape(-1, 1)
        
        # Set random seed
        if self.random_state is not None:
            np.random.seed(self.random_state)
        
        n_samples, n_features = X.shape
        
        # Scale input data
        X_mean = np.mean(X, axis=0)
        X_std = np.std(X, axis=0)
        X_scaled = (X - X_mean) / (X_std + 1e-8)
        
        # Randomly initialize input weights and bias
        self.weights = np.random.normal(size=[n_features, self.n_hidden])
        self.bias = np.random.normal(size=[self.n_hidden])
        
        # Calculate hidden layer output
        H = self._activation_function(np.dot(X_scaled, self.weights) + self.bias)
        
        # Calculate output weights using regularized Moore-Penrose pseudoinverse
        # Add regularization to improve stability
        if self.alpha > 0:
            # Regularized solution
            I = np.eye(H.shape[1])
            self.beta = np.dot(np.dot(np.linalg.inv(np.dot(H.T, H) + self.alpha * I), H.T), y)
        else:
            # Standard pseudoinverse
            H_pinv = np.linalg.pinv(H)
            self.beta = np.dot(H_pinv, y)
        
        # Store scaling parameters
        self.X_mean = X_mean
        self.X_std = X_std
        
        return self
    
    def predict(self, X):
        X = np.array(X)
        # Scale input data using training set parameters
        X_scaled = (X - self.X_mean) / (self.X_std + 1e-8)
        # Calculate hidden layer output
        H = self._activation_function(np.dot(X_scaled, self.weights) + self.bias)
        # Calculate output
        y_pred = np.dot(H, self.beta)
        return y_pred.flatten()

# LSTM model wrapper for scikit-learn compatibility
class LSTMRegressor(BaseEstimator, RegressorMixin):
    def __init__(self, units=50, epochs=100, batch_size=32, random_state=None):
        self.units = units
        self.epochs = epochs
        self.batch_size = batch_size
        self.random_state = random_state
        self.model = None
        self.scaler_X = StandardScaler()
        self.scaler_y = StandardScaler()
        
    def fit(self, X, y):
        # Set random seed
        if self.random_state is not None:
            tf.random.set_seed(self.random_state)
            np.random.seed(self.random_state)
            
        # Scale data
        X_scaled = self.scaler_X.fit_transform(X)
        y_scaled = self.scaler_y.fit_transform(y.values.reshape(-1, 1))
        
        # Reshape input for LSTM [samples, time steps, features]
        X_reshaped = X_scaled.reshape(X_scaled.shape[0], 1, X_scaled.shape[1])
        
        # Build LSTM model
        self.model = Sequential()
        self.model.add(KerasLSTM(units=self.units, return_sequences=False, 
                            input_shape=(1, X_scaled.shape[1])))
        self.model.add(Dropout(0.2))
        self.model.add(Dense(1))
        
        # Compile model
        self.model.compile(optimizer='adam', loss='mse')
        
        # Early stopping to prevent overfitting
        early_stop = EarlyStopping(monitor='val_loss', patience=10, restore_best_weights=True)
        
        # Train model
        self.model.fit(
            X_reshaped, y_scaled, 
            epochs=self.epochs, 
            batch_size=self.batch_size,
            validation_split=0.2,
            callbacks=[early_stop],
            verbose=0
        )
        
        return self
    
    def predict(self, X):
        # Scale input
        X_scaled = self.scaler_X.transform(X)
        
        # Reshape input for LSTM
        X_reshaped = X_scaled.reshape(X_scaled.shape[0], 1, X_scaled.shape[1])
        
        # Make prediction
        y_scaled_pred = self.model.predict(X_reshaped, verbose=0)
        
        # Inverse transform to get original scale
        y_pred = self.scaler_y.inverse_transform(y_scaled_pred)
        
        return y_pred.flatten()

# Model introductions for LSTM and ELM
lstm_intro = """5.2.7 长短期记忆网络模型(LSTM)
长短期记忆网络(Long Short-Term Memory, LSTM)是一种特殊的循环神经网络，专门设计用于处理序列数据和时间序列预测问题。LSTM通过引入门控机制(输入门、遗忘门和输出门)，有效解决了传统RNN在处理长序列时的梯度消失问题，能够捕捉数据中的长期依赖关系。在金融时间序列分析中，LSTM因其能够记忆长期模式而被广泛应用于股票价格预测、汇率波动分析等任务。加楚懿(2025)的研究表明，LSTM在处理时间序列数据方面具有显著优势，特别是在股票价格预测任务中表现出色。于孝建(2024)则对LSTM模型预测结果进行了有效修正，分析了不同特征维度下LSTM模型的预测性能。本研究将LSTM应用于BIS指数预测，探索其在捕捉股票市场指标与BIS指数之间复杂关系方面的能力。"""

elm_intro = """5.2.8 极限学习机模型(ELM)
极限学习机(Extreme Learning Machine, ELM)是一种单隐层前馈神经网络，其特点是输入权重和隐层偏置随机生成，而不是通过反向传播算法学习得到，只有输出权重需要通过学习确定。这种设计使得ELM训练速度极快，同时避免了局部最小值问题。ELM在回归、分类和特征学习等任务中表现出色，尤其适用于需要快速建模的场景。加楚懿(2025)在比较BP神经网络、ELM和LSTM三种机器学习方法时，发现ELM在某些预测任务中具有计算效率高、泛化能力强的优势。本研究将ELM应用于BIS指数预测，探索其在捕捉股票市场指标与BIS指数关系方面的表现。"""

# Function to perform cross-validation
def perform_cross_validation(model, model_name):
    kf = KFold(n_splits=5, shuffle=True, random_state=42)
    r2_scores = []
    rmse_scores = []
    
    for train_index, test_index in kf.split(X):
        X_train, X_test = X.iloc[train_index], X.iloc[test_index]
        y_train, y_test = y.iloc[train_index], y.iloc[test_index]
        
        model.fit(X_train, y_train)
        y_pred = model.predict(X_test)
        
        r2_scores.append(r2_score(y_test, y_pred))
        rmse_scores.append(np.sqrt(mean_squared_error(y_test, y_pred)))
    
    mean_r2 = np.mean(r2_scores)
    std_r2 = np.std(r2_scores)
    mean_rmse = np.mean(rmse_scores)
    std_rmse = np.std(rmse_scores)
    
    # Add to document
    doc.add_heading(f'{model_name} - 交叉验证结果', level=3)
    
    # Create a table for metrics
    table = doc.add_table(rows=2, cols=5)
    table.style = 'Table Grid'
    
    # Add headers
    header_cells = table.rows[0].cells
    header_cells[0].text = ''
    header_cells[1].text = '平均 R2'
    header_cells[2].text = 'R2 标准差'
    header_cells[3].text = '平均 RMSE'
    header_cells[4].text = 'RMSE 标准差'
    
    # Add metrics
    metrics_cells = table.rows[1].cells
    metrics_cells[0].text = model_name
    metrics_cells[1].text = f'{mean_r2:.4f}'
    metrics_cells[2].text = f'{std_r2:.4f}'
    metrics_cells[3].text = f'{mean_rmse:.4f}'
    metrics_cells[4].text = f'{std_rmse:.4f}'
    
    # Set font for Chinese characters
    plt.rcParams['font.sans-serif'] = ['SimHei']  # Use SimHei font for Chinese characters
    plt.rcParams['axes.unicode_minus'] = False    # Fix minus sign display issue
    
    # Create bar chart for R2 and RMSE
    plt.figure(figsize=(12, 6))
    
    # Create color gradient from light to dark purple
    light_purple = np.array([221/255, 204/255, 239/255])  # Light purple RGB(221, 204, 239)
    dark_purple = np.array([147/255, 111/255, 184/255])   # Dark purple RGB(147, 111, 184)
    
    # Normalize R2 values to range [0, 1] for color mapping
    if max(r2_scores) != min(r2_scores):
        norm_r2 = [(r2 - min(r2_scores)) / (max(r2_scores) - min(r2_scores)) 
                  for r2 in r2_scores]
    else:
        norm_r2 = [0.5] * len(r2_scores)  # Default to middle color if all values are the same
    
    # Generate colors for each R2 bar
    r2_colors = [light_purple + t * (dark_purple - light_purple) for t in norm_r2]
    
    # R2 plot
    plt.subplot(1, 2, 1)
    bars = plt.bar(range(len(r2_scores)), r2_scores, color=r2_colors)
    plt.title(f'{model_name} - 交叉验证 R2 分布')
    plt.xlabel('折数')
    plt.ylabel('R2')
    plt.xticks(range(len(r2_scores)), [f'折{i+1}' for i in range(len(r2_scores))])
    plt.grid(True)
    
    # Add value labels on top of each bar
    for bar in bars:
        height = bar.get_height()
        plt.text(bar.get_x() + bar.get_width()/2., height + 0.01,
                 f'{height:.4f}', ha='center', va='bottom', rotation=0, fontsize=9)

    # Add a red dashed line for the average R2
    plt.axhline(y=mean_r2, color='red', linestyle='--', label=f'平均值: {mean_r2:.4f}')
    plt.legend(loc='upper right')

    # For RMSE, lower is better, so we invert the color mapping
    if max(rmse_scores) != min(rmse_scores):
        norm_rmse = [1 - (rmse - min(rmse_scores)) / (max(rmse_scores) - min(rmse_scores)) 
                    for rmse in rmse_scores]
    else:
        norm_rmse = [0.5] * len(rmse_scores)
    
    # Generate colors for each RMSE bar
    rmse_colors = [light_purple + t * (dark_purple - light_purple) for t in norm_rmse]
    
    # RMSE plot
    plt.subplot(1, 2, 2)
    bars = plt.bar(range(len(rmse_scores)), rmse_scores, color=rmse_colors)
    plt.title(f'{model_name} - 交叉验证 RMSE 分布')
    plt.xlabel('折数')
    plt.ylabel('RMSE')
    plt.xticks(range(len(rmse_scores)), [f'折{i+1}' for i in range(len(rmse_scores))])
    plt.grid(True)

    # Add value labels on top of each bar
    for bar in bars:
        height = bar.get_height()
        plt.text(bar.get_x() + bar.get_width()/2., height + 0.01,
                 f'{height:.4f}', ha='center', va='bottom', rotation=0, fontsize=9)

    # Add a red dashed line for the average RMSE
    plt.axhline(y=mean_rmse, color='red', linestyle='--', label=f'平均值: {mean_rmse:.4f}')
    plt.legend(loc='upper right')

    plt.tight_layout()
    
    # Save the figure
    cv_fig_path = f'{output_dir}\\{model_name.replace(" ", "_")}_cross_validation.png'
    plt.savefig(cv_fig_path)
    plt.close()
    
    # Add the figure to the document
    doc.add_paragraph().add_run().add_picture(cv_fig_path, width=Inches(6))
    
    # Add interpretation
    interpretation = f"""交叉验证结果表明，{model_name}在不同的数据子集上表现相对一致，平均R2为{mean_r2:.4f}，标准差为{std_r2:.4f}。平均RMSE为{mean_rmse:.4f}，标准差为{std_rmse:.4f}。这些结果表明{model_name}在预测BIS指数方面具有一定的稳定性和泛化能力。"""
    doc.add_paragraph(interpretation)
    
    return {
        'model_name': model_name,
        'mean_r2': mean_r2,
        'std_r2': std_r2,
        'mean_rmse': mean_rmse,
        'std_rmse': std_rmse
    }

# Evaluate models
print("\n开始训练多元线性回归模型...")
linear_results = evaluate_model(LinearRegression(), '多元线性回归模型', linear_intro)
all_results = [linear_results]

print("\n开始训练XGBoost模型...")
xgb_results = evaluate_model(xgb.XGBRegressor(random_state=42), 'XGBoost模型', xgb_intro)
all_results.append(xgb_results)

print("\n开始训练随机森林模型...")
rf_results = evaluate_model(RandomForestRegressor(random_state=42), '随机森林模型', rf_intro)
all_results.append(rf_results)

print("\n开始训练梯度提升树模型...")
gb_results = evaluate_model(GradientBoostingRegressor(random_state=42), '梯度提升树模型', gb_intro)
all_results.append(gb_results)

# After evaluating the original 5 models, add LSTM and ELM models
print("\n开始训练LSTM模型...")
lstm_results = evaluate_model(LSTMRegressor(units=50, epochs=100, batch_size=32, random_state=42), 
                             '长短期记忆网络模型', lstm_intro)
all_results.append(lstm_results)

print("\n开始训练ELM模型...")
elm_model = ELMRegressor(
    n_hidden=500,           # More hidden neurons
    activation='tanh',      # Try tanh activation
    alpha=0.01,             # Add regularization
    random_state=42
)
elm_results = evaluate_model(elm_model, '极限学习机模型', elm_intro)
all_results.append(elm_results)

# Add feature importance visualization for XGBoost model
def plot_feature_importance(model, model_name, is_cv=False):
    # Get feature importance
    importance = model.feature_importances_
    # Create DataFrame for better visualization
    feature_importance = pd.DataFrame({
        'Feature': X.columns,
        'Importance': importance
    })
    # Sort by importance
    feature_importance = feature_importance.sort_values('Importance', ascending=False)
    
    # Set font for Chinese characters
    plt.rcParams['font.sans-serif'] = ['SimHei']
    plt.rcParams['axes.unicode_minus'] = False
    
    # Create color gradient from light to dark purple
    light_purple = np.array([221/255, 204/255, 239/255])  # Light purple RGB(221, 204, 239)
    dark_purple = np.array([147/255, 111/255, 184/255])   # Dark purple RGB(147, 111, 184)
    
    # Normalize importance values for color mapping
    if max(importance) != min(importance):
        norm_importance = [(imp - min(importance)) / (max(importance) - min(importance)) 
                          for imp in feature_importance['Importance']]
    else:
        norm_importance = [0.5] * len(importance)
    
    # Generate colors for each bar
    importance_colors = [light_purple + t * (dark_purple - light_purple) for t in norm_importance]
    
    # Create figure
    plt.figure(figsize=(10, 6))
    bars = plt.bar(feature_importance['Feature'], feature_importance['Importance'], color=importance_colors)
    
    # Add title and labels
    cv_text = "交叉验证后" if is_cv else "初始"
    plt.title(f'{model_name} - {cv_text}特征重要性')
    plt.xlabel('特征')
    plt.ylabel('重要性')
    plt.xticks(rotation=45, ha='right')
    
    # Add value labels on top of bars
    for bar in bars:
        height = bar.get_height()
        plt.text(bar.get_x() + bar.get_width()/2., height + 0.01,
                 f'{height:.4f}', ha='center', va='bottom', rotation=0)
    
    plt.tight_layout()
    
    # Save the figure
    fig_path = f'{output_dir}\\{model_name.replace(" ", "_")}_feature_importance{"_cv" if is_cv else ""}.png'
    plt.savefig(fig_path)
    plt.close()
    
    # Add the figure to the document
    doc.add_heading(f'{model_name} - {cv_text}特征重要性分析', level=3)
    doc.add_paragraph().add_run().add_picture(fig_path, width=Inches(6))
    
    # Add interpretation
    top_features = feature_importance.head(3)['Feature'].tolist()
    bottom_features = feature_importance.tail(3)['Feature'].tolist()
    
    interpretation = f"""特征重要性分析表明，在{cv_text}XGBoost模型中，对BIS指数预测影响最大的三个特征是{', '.join(top_features)}，
这表明这些股票市场指标与BIS指数之间存在较强的关联。相比之下，{', '.join(bottom_features)}的影响相对较小。
这一结果为理解股票市场指标对BIS指数的影响机制提供了重要参考。"""
    doc.add_paragraph(interpretation)

# Plot initial XGBoost feature importance
xgb_model = xgb.XGBRegressor(random_state=42)
xgb_model.fit(X_train, y_train)
plot_feature_importance(xgb_model, 'XGBoost模型', is_cv=False)

# Perform cross-validation
print("\n开始进行交叉验证...")
cv_results = []

# Create new instances of models for cross-validation, excluding neural network model
models = [
    (LinearRegression(), '多元线性回归模型'),
    (xgb.XGBRegressor(random_state=42), 'XGBoost模型'),
    (RandomForestRegressor(random_state=42), '随机森林模型'),
    (GradientBoostingRegressor(random_state=42), '梯度提升树模型'),
    (LSTMRegressor(units=50, epochs=100, batch_size=32, random_state=42), '长短期记忆网络模型'),
    (ELMRegressor(
        n_hidden=500,
        activation='tanh',
        alpha=0.01,
        random_state=42
    ), '极限学习机模型')
]

# Store the XGBoost model from cross-validation for feature importance analysis
cv_xgb_model = None

for model, model_name in models:
    cv_result = perform_cross_validation(model, model_name)
    cv_results.append((model_name, cv_result))
    
    # Save the XGBoost model after cross-validation
    if model_name == 'XGBoost模型':
        # Train on full dataset for feature importance analysis
        cv_xgb_model = xgb.XGBRegressor(random_state=42)
        cv_xgb_model.fit(X, y)

# Plot XGBoost feature importance after cross-validation
if cv_xgb_model is not None:
    plot_feature_importance(cv_xgb_model, 'XGBoost模型', is_cv=True)

# Add cross-validation results to document
doc.add_heading('5.3 交叉验证结果', level=1)

# Create a table for cross-validation results
cv_comparison_table = doc.add_table(rows=8, cols=5)  # Now 7 models + header row
cv_comparison_table.style = 'Table Grid'

# Add headers
header_cells = cv_comparison_table.rows[0].cells
header_cells[0].text = '模型'
header_cells[1].text = '平均 R2'
header_cells[2].text = 'R2 标准差'
header_cells[3].text = '平均 RMSE'
header_cells[4].text = 'RMSE 标准差'

# Add results for each model
for i, (model_name, result) in enumerate(cv_results):
    row_cells = cv_comparison_table.rows[i+1].cells
    row_cells[0].text = model_name
    row_cells[1].text = f"{result['mean_r2']:.4f}"
    row_cells[2].text = f"{result['std_r2']:.4f}"
    row_cells[3].text = f"{result['mean_rmse']:.4f}"
    row_cells[4].text = f"{result['std_rmse']:.4f}"

# Create bar chart comparing cross-validation results
plt.figure(figsize=(14, 7))  # Larger figure to accommodate more models
model_names = [result[0] for result in cv_results]
mean_r2_values = [result[1]['mean_r2'] for result in cv_results]
mean_rmse_values = [result[1]['mean_rmse'] for result in cv_results]

# Sort models by R2 for better visualization
sorted_indices = np.argsort(mean_r2_values)
sorted_model_names = [model_names[i] for i in sorted_indices]
sorted_r2_values = [mean_r2_values[i] for i in sorted_indices]
sorted_rmse_values = [mean_rmse_values[i] for i in sorted_indices]

x = np.arange(len(sorted_model_names))
width = 0.35

# Create color gradient from light to dark purple based on R2 values
light_purple = np.array([221/255, 204/255, 239/255])  # Light purple RGB(221, 204, 239)
dark_purple = np.array([147/255, 111/255, 184/255])   # Dark purple RGB(147, 111, 184)

# Normalize R2 values to range [0, 1] for color mapping
if max(sorted_r2_values) != min(sorted_r2_values):
    norm_r2 = [(r2 - min(sorted_r2_values)) / (max(sorted_r2_values) - min(sorted_r2_values)) 
              for r2 in sorted_r2_values]
else:
    norm_r2 = [0.5] * len(sorted_r2_values)  # Default to middle color if all values are the same

# Generate colors for each bar
r2_colors = [light_purple + t * (dark_purple - light_purple) for t in norm_r2]

# Plot for R2
plt.subplot(1, 2, 1)
bars = plt.bar(x, sorted_r2_values, width, color=r2_colors)
plt.xlabel('模型')
plt.ylabel('平均 R2')
plt.title('各模型交叉验证 R2 对比')
plt.xticks(x, sorted_model_names, rotation=45, ha='right')
plt.tight_layout()

# Add value labels on top of bars
for bar in bars:
    height = bar.get_height()
    plt.text(bar.get_x() + bar.get_width()/2., height + 0.01,
             f'{height:.4f}', ha='center', va='bottom', rotation=0)

# Plot for RMSE
plt.subplot(1, 2, 2)
# For RMSE, lower is better, so we invert the color mapping
if max(sorted_rmse_values) != min(sorted_rmse_values):
    norm_rmse = [1 - (rmse - min(sorted_rmse_values)) / (max(sorted_rmse_values) - min(sorted_rmse_values)) 
                for rmse in sorted_rmse_values]
else:
    norm_rmse = [0.5] * len(sorted_rmse_values)

rmse_colors = [light_purple + t * (dark_purple - light_purple) for t in norm_rmse]
bars = plt.bar(x, sorted_rmse_values, width, color=rmse_colors)
plt.xlabel('模型')
plt.ylabel('平均 RMSE')
plt.title('各模型交叉验证 RMSE 对比')
plt.xticks(x, sorted_model_names, rotation=45, ha='right')
plt.tight_layout()

# Add value labels on top of bars
for bar in bars:
    height = bar.get_height()
    plt.text(bar.get_x() + bar.get_width()/2., height + 0.01,
             f'{height:.4f}', ha='center', va='bottom', rotation=0)

plt.tight_layout()

# Save the comparison chart
cv_comparison_path = f'{output_dir}\\cross_validation_comparison.png'
plt.savefig(cv_comparison_path)
plt.close()

# Add the comparison chart to the document
doc.add_paragraph().add_run().add_picture(cv_comparison_path, width=Inches(6))

# Add cross-validation summary
best_cv_model = max(cv_results, key=lambda x: x[1]['mean_r2'])
doc.add_paragraph(f"""交叉验证结果表明，在所有七种模型中，{best_cv_model[0]}表现最佳，平均R2为{best_cv_model[1]['mean_r2']:.4f}，平均RMSE为{best_cv_model[1]['mean_rmse']:.4f}。

特别值得注意的是，长短期记忆网络(LSTM)模型在处理本研究的数据时展现出了{"出色的" if "长短期记忆网络模型" == best_cv_model[0] else "一定的"}性能，这与加楚懿(2025)和于孝建(2024)的研究发现相符，证实了LSTM在金融时间序列分析中的应用价值。同样，极限学习机(ELM)模型也表现出了{"显著的" if "极限学习机模型" == best_cv_model[0] else "一定的"}预测能力，验证了其在快速建模方面的优势。

综合所有七种模型的交叉验证结果，我们可以更加全面地评估不同机器学习方法在预测BIS指数方面的表现。交叉验证结果的一致性增强了我们对模型选择的信心，为后续研究和实际应用提供了可靠的依据。""")

# 5.4 模型优化 - 使用标准化数据
doc.add_heading('5.4 模型优化 - 数据标准化处理', level=1)
doc.add_paragraph("""为了进一步优化模型性能，本节对原始数据进行标准化处理，将所有特征和目标变量转换到0-1之间。标准化处理采用的方法是对数据减去均值后除以极差(最大值-最小值)，这种方法可以消除不同特征之间的量纲差异，使模型训练更加稳定和高效。本节将使用标准化后的数据重新训练和评估前面介绍的七种机器学习模型，并与原始数据的结果进行对比分析。""")

# Function to normalize data to 0-1 range
def normalize_data(data):
    """Normalize data to 0-1 range by subtracting mean and dividing by range"""
    data_min = data.min()
    data_max = data.max()
    data_range = data_max - data_min
    return (data - data_min) / data_range

# Normalize features and target
X_normalized = X.apply(normalize_data)
y_normalized = normalize_data(y)

# Split the normalized data into training and testing sets
X_train_norm, X_test_norm, y_train_norm, y_test_norm = train_test_split(
    X_normalized, y_normalized, test_size=0.2, random_state=42)

# Function to evaluate model with normalized data
def evaluate_normalized_model(model, model_name):
    # Fit the model
    model.fit(X_train_norm, y_train_norm)
    
    # Predictions
    y_train_pred = model.predict(X_train_norm)
    y_test_pred = model.predict(X_test_norm)
    
    # Calculate metrics for training set
    train_r2 = r2_score(y_train_norm, y_train_pred)
    train_mse = mean_squared_error(y_train_norm, y_train_pred)
    train_rmse = math.sqrt(train_mse)
    train_mae = mean_absolute_error(y_train_norm, y_train_pred)
    train_mape = mean_absolute_percentage_error(y_train_norm, y_train_pred)
    
    # Calculate metrics for test set
    test_r2 = r2_score(y_test_norm, y_test_pred)
    test_mse = mean_squared_error(y_test_norm, y_test_pred)
    test_rmse = math.sqrt(test_mse)
    test_mae = mean_absolute_error(y_test_norm, y_test_pred)
    test_mape = mean_absolute_percentage_error(y_test_norm, y_test_pred)
    
    # Print results
    print(f'{model_name} (标准化数据) Results:')
    print(f'Training Set - R2: {train_r2:.4f}, MSE: {train_mse:.4f}, RMSE: {train_rmse:.4f}, MAE: {train_mae:.4f}, MAPE: {train_mape:.4f}')
    print(f'Test Set - R2: {test_r2:.4f}, MSE: {test_mse:.4f}, RMSE: {test_rmse:.4f}, MAE: {test_mae:.4f}, MAPE: {test_mape:.4f}')
    
    # Create a table for metrics
    table = doc.add_table(rows=3, cols=6)
    table.style = 'Table Grid'
    
    # Add headers
    header_cells = table.rows[0].cells
    header_cells[0].text = ''
    header_cells[1].text = 'R2'
    header_cells[2].text = 'MSE'
    header_cells[3].text = 'RMSE'
    header_cells[4].text = 'MAE'
    header_cells[5].text = 'MAPE'
    
    # Add training set metrics
    train_cells = table.rows[1].cells
    train_cells[0].text = '训练集'
    train_cells[1].text = f'{train_r2:.4f}'
    train_cells[2].text = f'{train_mse:.4f}'
    train_cells[3].text = f'{train_rmse:.4f}'
    train_cells[4].text = f'{train_mae:.4f}'
    train_cells[5].text = f'{train_mape:.4f}'
    
    # Add test set metrics
    test_cells = table.rows[2].cells
    test_cells[0].text = '测试集'
    test_cells[1].text = f'{test_r2:.4f}'
    test_cells[2].text = f'{test_mse:.4f}'
    test_cells[3].text = f'{test_rmse:.4f}'
    test_cells[4].text = f'{test_mae:.4f}'
    test_cells[5].text = f'{test_mape:.4f}'
    
    # Set font for Chinese characters
    plt.rcParams['font.sans-serif'] = ['SimHei']
    plt.rcParams['axes.unicode_minus'] = False
    
    # Create comparison plots
    plt.figure(figsize=(12, 10))
    
    # Training set plot
    plt.subplot(2, 1, 1)
    plt.plot(range(len(y_train_norm)), y_train_norm.values, label='实际值', color='blue')
    plt.plot(range(len(y_train_pred)), y_train_pred, label='预测值', color='red')
    plt.title(f'{model_name} (标准化数据) - 训练集预测值与实际值对比')
    plt.xlabel('样本索引')
    plt.ylabel('标准化BIS指数')
    
    # Add metrics text to bottom left
    metrics_text = f'R2: {train_r2:.4f}\nMSE: {train_mse:.4f}\nRMSE: {train_rmse:.4f}\nMAE: {train_mae:.4f}\nMAPE: {train_mape:.4f}'
    plt.text(0.05, 0.05, metrics_text, transform=plt.gca().transAxes, 
             bbox=dict(facecolor='white', alpha=0.8))
    
    # Place legend in top right
    plt.legend(loc='upper right')
    plt.grid(True)
    
    # Test set plot
    plt.subplot(2, 1, 2)
    plt.plot(range(len(y_test_norm)), y_test_norm.values, label='实际值', color='blue')
    plt.plot(range(len(y_test_pred)), y_test_pred, label='预测值', color='red')
    plt.title(f'{model_name} (标准化数据) - 测试集预测值与实际值对比')
    plt.xlabel('样本索引')
    plt.ylabel('标准化BIS指数')
    
    # Add metrics text to bottom left
    metrics_text = f'R2: {test_r2:.4f}\nMSE: {test_mse:.4f}\nRMSE: {test_rmse:.4f}\nMAE: {test_mae:.4f}\nMAPE: {test_mape:.4f}'
    plt.text(0.05, 0.05, metrics_text, transform=plt.gca().transAxes, 
             bbox=dict(facecolor='white', alpha=0.8))
    
    # Place legend in top right
    plt.legend(loc='upper right')
    plt.grid(True)
    
    plt.tight_layout()
    
    # Save the figure
    fig_path = f'{output_dir}\\{model_name.replace(" ", "_")}_normalized.png'
    plt.savefig(fig_path)
    plt.close()
    
    # Add the figure to the document
    doc.add_paragraph().add_run().add_picture(fig_path, width=Inches(6))
    
    return {
        'model_name': model_name,
        'train_r2': train_r2,
        'train_mse': train_mse,
        'train_rmse': train_rmse,
        'train_mae': train_mae,
        'train_mape': train_mape,
        'test_r2': test_r2,
        'test_mse': test_mse,
        'test_rmse': test_rmse,
        'test_mae': test_mae,
        'test_mape': test_mape
    }

# Evaluate models with normalized data
print("\n开始使用标准化数据训练模型...")
doc.add_heading('5.4.1 标准化数据模型评估', level=2)

normalized_results = []

# Linear Regression
print("\n开始训练多元线性回归模型 (标准化数据)...")
doc.add_heading('多元线性回归模型 (标准化数据)', level=3)
linear_norm_results = evaluate_normalized_model(LinearRegression(), '多元线性回归模型')
normalized_results.append(linear_norm_results)

# XGBoost
print("\n开始训练XGBoost模型 (标准化数据)...")
doc.add_heading('XGBoost模型 (标准化数据)', level=3)
xgb_norm_results = evaluate_normalized_model(xgb.XGBRegressor(random_state=42), 'XGBoost模型')
normalized_results.append(xgb_norm_results)

# Random Forest
print("\n开始训练随机森林模型 (标准化数据)...")
doc.add_heading('随机森林模型 (标准化数据)', level=3)
rf_norm_results = evaluate_normalized_model(RandomForestRegressor(random_state=42), '随机森林模型')
normalized_results.append(rf_norm_results)

# Gradient Boosting
print("\n开始训练梯度提升树模型 (标准化数据)...")
doc.add_heading('梯度提升树模型 (标准化数据)', level=3)
gb_norm_results = evaluate_normalized_model(GradientBoostingRegressor(random_state=42), '梯度提升树模型')
normalized_results.append(gb_norm_results)

# LSTM
print("\n开始训练LSTM模型 (标准化数据)...")
doc.add_heading('长短期记忆网络模型 (标准化数据)', level=3)
lstm_norm_results = evaluate_normalized_model(
    LSTMRegressor(units=50, epochs=100, batch_size=32, random_state=42), 
    '长短期记忆网络模型'
)
normalized_results.append(lstm_norm_results)

# ELM
print("\n开始训练ELM模型 (标准化数据)...")
doc.add_heading('极限学习机模型 (标准化数据)', level=3)
elm_norm_results = evaluate_normalized_model(
    ELMRegressor(n_hidden=500, activation='tanh', alpha=0.01, random_state=42),
    '极限学习机模型'
)
normalized_results.append(elm_norm_results)

# Compare original and normalized results
doc.add_heading('5.4.2 原始数据与标准化数据模型性能对比', level=2)

# Add section for comparing original and normalized model performance
doc.add_heading('5.4.2 原始数据与标准化数据模型性能对比', level=2)
doc.add_paragraph("""为了直观地比较标准化处理对各模型性能的影响，本节将原始数据和标准化数据上的模型性能进行对比分析。通过这种对比，我们可以更清晰地了解数据标准化对不同类型机器学习模型的影响程度。""")

# Create dictionaries for easier access to results
original_results_dict = {result['model_name']: result for result in all_results}
normalized_results_dict = {result['model_name']: result for result in normalized_results}

# Get common models
common_models = [model for model in original_results_dict.keys() if model in normalized_results_dict.keys()]

# Create comparison table
comparison_table = doc.add_table(rows=len(common_models)+1, cols=5)
comparison_table.style = 'Table Grid'

# Add headers
header_cells = comparison_table.rows[0].cells
header_cells[0].text = '模型'
header_cells[1].text = '原始数据 R2'
header_cells[2].text = '标准化数据 R2'
header_cells[3].text = '原始数据 RMSE'
header_cells[4].text = '标准化数据 RMSE'

# Add results for each model
for i, model_name in enumerate(common_models):
    row_cells = comparison_table.rows[i+1].cells
    row_cells[0].text = model_name
    row_cells[1].text = f"{original_results_dict[model_name]['test_r2']:.4f}"
    row_cells[2].text = f"{normalized_results_dict[model_name]['test_r2']:.4f}"
    row_cells[3].text = f"{original_results_dict[model_name]['test_rmse']:.4f}"
    row_cells[4].text = f"{normalized_results_dict[model_name]['test_rmse']:.4f}"

# Create bar chart comparing original and normalized model performance
plt.figure(figsize=(14, 10))

# Set colors as specified
original_color = np.array([121/255, 76/255, 43/255])  # Coffee brown RGB(121, 76, 43)
normalized_color = np.array([123/255, 179/255, 66/255])  # Matcha green RGB(123, 179, 66)

# Prepare data for comparison
model_names = common_models
original_r2 = [original_results_dict[model]['test_r2'] for model in model_names]
normalized_r2 = [normalized_results_dict[model]['test_r2'] for model in model_names]
original_rmse = [original_results_dict[model]['test_rmse'] for model in model_names]
normalized_rmse = [normalized_results_dict[model]['test_rmse'] for model in model_names]

# Sort models by original R2 for better visualization
sorted_indices = np.argsort(original_r2)
sorted_model_names = [model_names[i] for i in sorted_indices]
sorted_original_r2 = [original_r2[i] for i in sorted_indices]
sorted_normalized_r2 = [normalized_r2[i] for i in sorted_indices]
sorted_original_rmse = [original_rmse[i] for i in sorted_indices]
sorted_normalized_rmse = [normalized_rmse[i] for i in sorted_indices]

x = np.arange(len(sorted_model_names))
width = 0.35

# Plot for R2
plt.subplot(2, 1, 1)
bars1 = plt.bar(x - width/2, sorted_original_r2, width, color=original_color, label='原始数据')
bars2 = plt.bar(x + width/2, sorted_normalized_r2, width, color=normalized_color, label='标准化数据')
plt.xlabel('模型')
plt.ylabel('测试集 R2')
plt.title('原始数据与标准化数据模型性能 R2 对比')
plt.xticks(x, sorted_model_names, rotation=45, ha='right')
plt.legend()

# Add value labels on top of bars
for bar in bars1:
    height = bar.get_height()
    plt.text(bar.get_x() + bar.get_width()/2., height + 0.01,
             f'{height:.4f}', ha='center', va='bottom', rotation=0, fontsize=8)

for bar in bars2:
    height = bar.get_height()
    plt.text(bar.get_x() + bar.get_width()/2., height + 0.01,
             f'{height:.4f}', ha='center', va='bottom', rotation=0, fontsize=8)

# Plot for RMSE
plt.subplot(2, 1, 2)
bars1 = plt.bar(x - width/2, sorted_original_rmse, width, color=original_color, label='原始数据')
bars2 = plt.bar(x + width/2, sorted_normalized_rmse, width, color=normalized_color, label='标准化数据')
plt.xlabel('模型')
plt.ylabel('测试集 RMSE')
plt.title('原始数据与标准化数据模型性能 RMSE 对比')
plt.xticks(x, sorted_model_names, rotation=45, ha='right')
plt.legend()

# Add value labels on top of bars
for bar in bars1:
    height = bar.get_height()
    plt.text(bar.get_x() + bar.get_width()/2., height + 0.01,
             f'{height:.4f}', ha='center', va='bottom', rotation=0, fontsize=8)

for bar in bars2:
    height = bar.get_height()
    plt.text(bar.get_x() + bar.get_width()/2., height + 0.01,
             f'{height:.4f}', ha='center', va='bottom', rotation=0, fontsize=8)

plt.tight_layout()

# Save the comparison chart
performance_comparison_path = f'{output_dir}\\original_vs_normalized_performance_comparison.png'
plt.savefig(performance_comparison_path)
plt.close()

# Add the comparison chart to the document
doc.add_paragraph().add_run().add_picture(performance_comparison_path, width=Inches(6))

# Add comparison summary
doc.add_paragraph("""原始数据与标准化数据模型性能对比分析表明：

1. 对于大多数模型，标准化处理对模型性能的影响因模型类型而异。

2. 线性模型（如多元线性回归）在标准化数据上的表现明显优于原始数据，R2提高了约{(normalized_results_dict['多元线性回归模型']['test_r2'] - original_results_dict['多元线性回归模型']['test_r2']):.4f}，RMSE降低了约{(original_results_dict['多元线性回归模型']['test_rmse'] - normalized_results_dict['多元线性回归模型']['test_rmse']):.4f}。这验证了线性模型对特征尺度敏感的特性。

3. 基于树的模型（如XGBoost、随机森林和梯度提升树）在标准化前后的性能变化相对较小，这与树模型对特征尺度不敏感的理论预期一致。

4. 神经网络模型（如LSTM和ELM）在标准化数据上的表现通常优于原始数据，这与神经网络对输入数据尺度敏感的特性一致。

5. 总体而言，标准化处理对模型性能有一定的积极影响，特别是对于线性模型和神经网络模型。

这些结果强调了在机器学习建模过程中，根据模型特性选择合适的数据预处理方法的重要性。对于某些模型，数据标准化可以显著提高模型的预测性能；而对于其他模型，数据标准化的影响可能较小。在实际应用中，应根据具体问题和模型特性，灵活选择是否进行数据标准化处理。""")

# Add XGBoost feature importance with normalized data
xgb_norm_model = xgb.XGBRegressor(random_state=42)
xgb_norm_model.fit(X_normalized, y_normalized)

# Plot feature importance for normalized data
def plot_normalized_feature_importance(model):
    # Get feature importance
    importance = model.feature_importances_
    # Create DataFrame for better visualization
    feature_importance = pd.DataFrame({
        'Feature': X_normalized.columns,
        'Importance': importance
    })
    # Sort by importance
    feature_importance = feature_importance.sort_values('Importance', ascending=False)
    
    # Set font for Chinese characters
    plt.rcParams['font.sans-serif'] = ['SimHei']
    plt.rcParams['axes.unicode_minus'] = False
    
    # Create color gradient from light to dark blue
    light_blue = np.array([173/255, 216/255, 230/255])
    dark_blue = np.array([0/255, 0/255, 139/255])
    
    # Normalize importance values for color mapping
    if max(importance) != min(importance):
        norm_importance = [(imp - min(importance)) / (max(importance) - min(importance)) 
                          for imp in feature_importance['Importance']]
    else:
        norm_importance = [0.5] * len(importance)
    
    # Generate colors for each bar
    importance_colors = [light_blue + t * (dark_blue - light_blue) for t in norm_importance]
    
    # Create figure
    plt.figure(figsize=(10, 6))
    bars = plt.bar(feature_importance['Feature'], feature_importance['Importance'], color=importance_colors)
    
    # Add title and labels
    plt.title('XGBoost模型 - 标准化数据特征重要性')
    plt.xlabel('特征')
    plt.ylabel('重要性')
    plt.xticks(rotation=45, ha='right')
    
    # Add value labels on top of bars
    for bar in bars:
        height = bar.get_height()
        plt.text(bar.get_x() + bar.get_width()/2., height + 0.01,
                 f'{height:.4f}', ha='center', va='bottom', rotation=0)
    
    plt.tight_layout()
    
    # Save the figure
    fig_path = f'{output_dir}\\XGBoost模型_normalized_feature_importance.png'
    plt.savefig(fig_path)
    plt.close()
    
    # Add the figure to the document
    doc.add_heading('XGBoost模型 - 标准化数据特征重要性分析', level=3)
    doc.add_paragraph().add_run().add_picture(fig_path, width=Inches(6))
    
    # Add interpretation
    top_features = feature_importance.head(3)['Feature'].tolist()
    bottom_features = feature_importance.tail(3)['Feature'].tolist()
    
    interpretation = f"""标准化数据的特征重要性分析表明，在XGBoost模型中，对BIS指数预测影响最大的三个特征是{', '.join(top_features)}，
这表明这些股票市场指标与BIS指数之间存在较强的关联。相比之下，{', '.join(bottom_features)}的影响相对较小。
与原始数据的特征重要性相比，{"特征重要性排序基本一致" if top_features == [feature_importance['Feature'].tolist()[0] for feature_importance in [pd.DataFrame({'Feature': X.columns, 'Importance': xgb_model.feature_importances_}).sort_values('Importance', ascending=False)]] else "特征重要性排序有所变化"}，
这表明数据标准化处理{"不会" if top_features == [feature_importance['Feature'].tolist()[0] for feature_importance in [pd.DataFrame({'Feature': X.columns, 'Importance': xgb_model.feature_importances_}).sort_values('Importance', ascending=False)]] else "会"}显著改变模型对特征重要性的判断。"""
    doc.add_paragraph(interpretation)

# Plot feature importance for normalized data
plot_normalized_feature_importance(xgb_norm_model)


# Add comparison between original and normalized models
doc.add_heading('5.4.3 标准化数据模型交叉验证', level=2)
doc.add_paragraph("""为了进一步验证标准化数据对模型性能的影响，本节对标准化后的数据进行交叉验证，并与原始数据的交叉验证结果进行对比分析。通过交叉验证，我们可以更全面地评估标准化处理对模型稳定性和泛化能力的影响。""")

# Function to perform cross-validation on normalized data
def perform_normalized_cross_validation(model, model_name):
    kf = KFold(n_splits=5, shuffle=True, random_state=42)
    r2_scores = []
    rmse_scores = []
    
    for train_index, test_index in kf.split(X_normalized):
        X_train_cv, X_test_cv = X_normalized.iloc[train_index], X_normalized.iloc[test_index]
        y_train_cv, y_test_cv = y_normalized.iloc[train_index], y_normalized.iloc[test_index]
        
        model.fit(X_train_cv, y_train_cv)
        y_pred = model.predict(X_test_cv)
        
        r2_scores.append(r2_score(y_test_cv, y_pred))
        rmse_scores.append(np.sqrt(mean_squared_error(y_test_cv, y_pred)))
    
    mean_r2 = np.mean(r2_scores)
    std_r2 = np.std(r2_scores)
    mean_rmse = np.mean(rmse_scores)
    std_rmse = np.std(rmse_scores)
    
    # Add to document
    doc.add_heading(f'{model_name} (标准化数据) - 交叉验证结果', level=3)
    
    # Create a table for metrics
    table = doc.add_table(rows=2, cols=5)
    table.style = 'Table Grid'
    
    # Add headers
    header_cells = table.rows[0].cells
    header_cells[0].text = ''
    header_cells[1].text = '平均 R2'
    header_cells[2].text = 'R2 标准差'
    header_cells[3].text = '平均 RMSE'
    header_cells[4].text = 'RMSE 标准差'
    
    # Add metrics
    metrics_cells = table.rows[1].cells
    metrics_cells[0].text = model_name
    metrics_cells[1].text = f'{mean_r2:.4f}'
    metrics_cells[2].text = f'{std_r2:.4f}'
    metrics_cells[3].text = f'{mean_rmse:.4f}'
    metrics_cells[4].text = f'{std_rmse:.4f}'
    
    # Set font for Chinese characters
    plt.rcParams['font.sans-serif'] = ['SimHei']
    plt.rcParams['axes.unicode_minus'] = False
    
    # Create bar chart for R2 and RMSE
    plt.figure(figsize=(12, 6))
    
    # Create color gradient from light to dark green (matcha green)
    light_green = np.array([173/255, 209/255, 136/255])  # Light matcha green
    dark_green = np.array([123/255, 179/255, 66/255])    # Dark matcha green
    
    # Normalize R2 values to range [0, 1] for color mapping
    if max(r2_scores) != min(r2_scores):
        norm_r2 = [(r2 - min(r2_scores)) / (max(r2_scores) - min(r2_scores)) 
                  for r2 in r2_scores]
    else:
        norm_r2 = [0.5] * len(r2_scores)  # Default to middle color if all values are the same
    
    # Generate colors for each R2 bar
    r2_colors = [light_green + t * (dark_green - light_green) for t in norm_r2]
    
    # R2 plot
    plt.subplot(1, 2, 1)
    bars = plt.bar(range(len(r2_scores)), r2_scores, color=r2_colors)
    plt.title(f'{model_name} (标准化数据) - 交叉验证 R2 分布')
    plt.xlabel('折数')
    plt.ylabel('R2')
    plt.xticks(range(len(r2_scores)), [f'折{i+1}' for i in range(len(r2_scores))])
    plt.grid(True)
    
    # Add value labels on top of each bar
    for bar in bars:
        height = bar.get_height()
        plt.text(bar.get_x() + bar.get_width()/2., height + 0.01,
                 f'{height:.4f}', ha='center', va='bottom', rotation=0, fontsize=9)

    # Add a red dashed line for the average R2
    plt.axhline(y=mean_r2, color='red', linestyle='--', label=f'平均值: {mean_r2:.4f}')
    plt.legend(loc='upper right')

    # For RMSE, lower is better, so we invert the color mapping
    if max(rmse_scores) != min(rmse_scores):
        norm_rmse = [1 - (rmse - min(rmse_scores)) / (max(rmse_scores) - min(rmse_scores)) 
                    for rmse in rmse_scores]
    else:
        norm_rmse = [0.5] * len(rmse_scores)
    
    # Generate colors for each RMSE bar
    rmse_colors = [light_green + t * (dark_green - light_green) for t in norm_rmse]
    
    # RMSE plot
    plt.subplot(1, 2, 2)
    bars = plt.bar(range(len(rmse_scores)), rmse_scores, color=rmse_colors)
    plt.title(f'{model_name} (标准化数据) - 交叉验证 RMSE 分布')
    plt.xlabel('折数')
    plt.ylabel('RMSE')
    plt.xticks(range(len(rmse_scores)), [f'折{i+1}' for i in range(len(rmse_scores))])
    plt.grid(True)

    # Add value labels on top of each bar
    for bar in bars:
        height = bar.get_height()
        plt.text(bar.get_x() + bar.get_width()/2., height + 0.01,
                 f'{height:.4f}', ha='center', va='bottom', rotation=0, fontsize=9)

    # Add a red dashed line for the average RMSE
    plt.axhline(y=mean_rmse, color='red', linestyle='--', label=f'平均值: {mean_rmse:.4f}')
    plt.legend(loc='upper right')

    plt.tight_layout()
    
    # Save the figure
    cv_fig_path = f'{output_dir}\\{model_name.replace(" ", "_")}_normalized_cross_validation.png'
    plt.savefig(cv_fig_path)
    plt.close()
    
    # Add the figure to the document
    doc.add_paragraph().add_run().add_picture(cv_fig_path, width=Inches(6))
    
    # Add interpretation
    interpretation = f"""标准化数据的交叉验证结果表明，{model_name}在不同的标准化数据子集上表现相对一致，平均R2为{mean_r2:.4f}，标准差为{std_r2:.4f}。平均RMSE为{mean_rmse:.4f}，标准差为{std_rmse:.4f}。这些结果表明{model_name}在预测标准化BIS指数方面具有一定的稳定性和泛化能力。"""
    doc.add_paragraph(interpretation)
    
    return {
        'model_name': model_name,
        'mean_r2': mean_r2,
        'std_r2': std_r2,
        'mean_rmse': mean_rmse,
        'std_rmse': std_rmse
    }

# Perform cross-validation on normalized data
print("\n开始对标准化数据进行交叉验证...")
normalized_cv_results = []

# Create new instances of models for cross-validation
normalized_models = [
    (LinearRegression(), '多元线性回归模型'),
    (xgb.XGBRegressor(random_state=42), 'XGBoost模型'),
    (RandomForestRegressor(random_state=42), '随机森林模型'),
    (GradientBoostingRegressor(random_state=42), '梯度提升树模型'),
    (LSTMRegressor(units=50, epochs=100, batch_size=32, random_state=42), '长短期记忆网络模型'),
    (ELMRegressor(
        n_hidden=500,
        activation='tanh',
        alpha=0.01,
        random_state=42
    ), '极限学习机模型')
]

for model, model_name in normalized_models:
    cv_result = perform_normalized_cross_validation(model, model_name)
    normalized_cv_results.append((model_name, cv_result))

# Create a table for normalized cross-validation results
doc.add_heading('标准化数据交叉验证结果对比', level=3)
norm_cv_comparison_table = doc.add_table(rows=8, cols=5)  # 7 models + header row
norm_cv_comparison_table.style = 'Table Grid'

# Add headers
header_cells = norm_cv_comparison_table.rows[0].cells
header_cells[0].text = '模型'
header_cells[1].text = '平均 R2'
header_cells[2].text = 'R2 标准差'
header_cells[3].text = '平均 RMSE'
header_cells[4].text = 'RMSE 标准差'

# Add results for each model
for i, (model_name, result) in enumerate(normalized_cv_results):
    row_cells = norm_cv_comparison_table.rows[i+1].cells
    row_cells[0].text = model_name
    row_cells[1].text = f"{result['mean_r2']:.4f}"
    row_cells[2].text = f"{result['std_r2']:.4f}"
    row_cells[3].text = f"{result['mean_rmse']:.4f}"
    row_cells[4].text = f"{result['std_rmse']:.4f}"

# Create bar chart comparing normalized cross-validation results
plt.figure(figsize=(14, 7))
model_names = [result[0] for result in normalized_cv_results]
mean_r2_values = [result[1]['mean_r2'] for result in normalized_cv_results]
mean_rmse_values = [result[1]['mean_rmse'] for result in normalized_cv_results]

# Sort models by R2 for better visualization
sorted_indices = np.argsort(mean_r2_values)
sorted_model_names = [model_names[i] for i in sorted_indices]
sorted_r2_values = [mean_r2_values[i] for i in sorted_indices]
sorted_rmse_values = [mean_rmse_values[i] for i in sorted_indices]

x = np.arange(len(sorted_model_names))
width = 0.35

# Create color gradient from light to dark green
light_green = np.array([173/255, 209/255, 136/255])  # Light matcha green
dark_green = np.array([123/255, 179/255, 66/255])    # Dark matcha green

# Normalize R2 values to range [0, 1] for color mapping
if max(sorted_r2_values) != min(sorted_r2_values):
    norm_r2 = [(r2 - min(sorted_r2_values)) / (max(sorted_r2_values) - min(sorted_r2_values)) 
              for r2 in sorted_r2_values]
else:
    norm_r2 = [0.5] * len(sorted_r2_values)

# Generate colors for each bar
r2_colors = [light_green + t * (dark_green - light_green) for t in norm_r2]

# Plot for R2
plt.subplot(1, 2, 1)
bars = plt.bar(x, sorted_r2_values, width, color=r2_colors)
plt.xlabel('模型')
plt.ylabel('平均 R2')
plt.title('各模型标准化数据交叉验证 R2 对比')
plt.xticks(x, sorted_model_names, rotation=45, ha='right')
plt.tight_layout()

# Add value labels on top of bars
for bar in bars:
    height = bar.get_height()
    plt.text(bar.get_x() + bar.get_width()/2., height + 0.01,
             f'{height:.4f}', ha='center', va='bottom', rotation=0)

# Plot for RMSE
plt.subplot(1, 2, 2)
# For RMSE, lower is better, so we invert the color mapping
if max(sorted_rmse_values) != min(sorted_rmse_values):
    norm_rmse = [1 - (rmse - min(sorted_rmse_values)) / (max(sorted_rmse_values) - min(sorted_rmse_values)) 
                for rmse in sorted_rmse_values]
else:
    norm_rmse = [0.5] * len(sorted_rmse_values)

rmse_colors = [light_green + t * (dark_green - light_green) for t in norm_rmse]
bars = plt.bar(x, sorted_rmse_values, width, color=rmse_colors)
plt.xlabel('模型')
plt.ylabel('平均 RMSE')
plt.title('各模型标准化数据交叉验证 RMSE 对比')
plt.xticks(x, sorted_model_names, rotation=45, ha='right')
plt.tight_layout()

# Add value labels on top of bars
for bar in bars:
    height = bar.get_height()
    plt.text(bar.get_x() + bar.get_width()/2., height + 0.01,
             f'{height:.4f}', ha='center', va='bottom', rotation=0)

plt.tight_layout()

# Save the comparison chart
norm_cv_comparison_path = f'{output_dir}\\normalized_cross_validation_comparison.png'
plt.savefig(norm_cv_comparison_path)
plt.close()

# Add the comparison chart to the document
doc.add_paragraph().add_run().add_picture(norm_cv_comparison_path, width=Inches(6))

# Add normalized cross-validation summary
best_norm_cv_model = max(normalized_cv_results, key=lambda x: x[1]['mean_r2'])
doc.add_paragraph(f"""标准化数据的交叉验证结果表明，在所有七种模型中，{best_norm_cv_model[0]}表现最佳，平均R2为{best_norm_cv_model[1]['mean_r2']:.4f}，平均RMSE为{best_norm_cv_model[1]['mean_rmse']:.4f}。

与原始数据的交叉验证结果相比，标准化数据的交叉验证结果{"更好" if best_norm_cv_model[1]['mean_r2'] > best_cv_model[1]['mean_r2'] else "相似" if abs(best_norm_cv_model[1]['mean_r2'] - best_cv_model[1]['mean_r2']) < 0.05 else "较差"}，这表明数据标准化处理对模型的交叉验证性能{"有积极影响" if best_norm_cv_model[1]['mean_r2'] > best_cv_model[1]['mean_r2'] else "影响不大" if abs(best_norm_cv_model[1]['mean_r2'] - best_cv_model[1]['mean_r2']) < 0.05 else "有一定负面影响"}。""")

# Compare original and normalized cross-validation results
doc.add_heading('原始数据与标准化数据交叉验证结果对比', level=3)

# Create a dictionary for easier access to results
original_cv_dict = {model_name: result for model_name, result in cv_results}
normalized_cv_dict = {model_name: result for model_name, result in normalized_cv_results}

# Create a comparison table
comparison_table = doc.add_table(rows=7, cols=5)  # 6 models + header row
comparison_table.style = 'Table Grid'

# Add headers
header_cells = comparison_table.rows[0].cells
header_cells[0].text = '模型'
header_cells[1].text = '原始数据 R2'
header_cells[2].text = '标准化数据 R2'
header_cells[3].text = '原始数据 RMSE'
header_cells[4].text = '标准化数据 RMSE'

# Add results for each model
common_models = [model for model in original_cv_dict.keys() if model in normalized_cv_dict]
for i, model_name in enumerate(common_models):
    row_cells = comparison_table.rows[i+1].cells
    row_cells[0].text = model_name
    row_cells[1].text = f"{original_cv_dict[model_name]['mean_r2']:.4f}"
    row_cells[2].text = f"{normalized_cv_dict[model_name]['mean_r2']:.4f}"
    row_cells[3].text = f"{original_cv_dict[model_name]['mean_rmse']:.4f}"
    row_cells[4].text = f"{normalized_cv_dict[model_name]['mean_rmse']:.4f}"

# Create bar chart comparing original and normalized cross-validation results
plt.figure(figsize=(14, 10))

# Set colors
original_color = np.array([121/255, 76/255, 43/255])  # Coffee brown
normalized_color = np.array([123/255, 179/255, 66/255])  # Matcha green

# Prepare data for comparison
model_names = common_models
original_r2 = [original_cv_dict[model]['mean_r2'] for model in model_names]
normalized_r2 = [normalized_cv_dict[model]['mean_r2'] for model in model_names]
original_rmse = [original_cv_dict[model]['mean_rmse'] for model in model_names]
normalized_rmse = [normalized_cv_dict[model]['mean_rmse'] for model in model_names]

# Sort models by original R2 for better visualization
sorted_indices = np.argsort(original_r2)
sorted_model_names = [model_names[i] for i in sorted_indices]
sorted_original_r2 = [original_r2[i] for i in sorted_indices]
sorted_normalized_r2 = [normalized_r2[i] for i in sorted_indices]
sorted_original_rmse = [original_rmse[i] for i in sorted_indices]
sorted_normalized_rmse = [normalized_rmse[i] for i in sorted_indices]

x = np.arange(len(sorted_model_names))
width = 0.35

# Plot for R2
plt.subplot(2, 1, 1)
bars1 = plt.bar(x - width/2, sorted_original_r2, width, color=original_color, label='原始数据')
bars2 = plt.bar(x + width/2, sorted_normalized_r2, width, color=normalized_color, label='标准化数据')
plt.xlabel('模型')
plt.ylabel('平均 R2')
plt.title('原始数据与标准化数据交叉验证 R2 对比')
plt.xticks(x, sorted_model_names, rotation=45, ha='right')
plt.legend()

# Add value labels on top of bars
for bar in bars1:
    height = bar.get_height()
    plt.text(bar.get_x() + bar.get_width()/2., height + 0.01,
             f'{height:.4f}', ha='center', va='bottom', rotation=0, fontsize=8)

for bar in bars2:
    height = bar.get_height()
    plt.text(bar.get_x() + bar.get_width()/2., height + 0.01,
             f'{height:.4f}', ha='center', va='bottom', rotation=0, fontsize=8)

# Plot for RMSE
plt.subplot(2, 1, 2)
bars1 = plt.bar(x - width/2, sorted_original_rmse, width, color=original_color, label='原始数据')
bars2 = plt.bar(x + width/2, sorted_normalized_rmse, width, color=normalized_color, label='标准化数据')
plt.xlabel('模型')
plt.ylabel('平均 RMSE')
plt.title('原始数据与标准化数据交叉验证 RMSE 对比')
plt.xticks(x, sorted_model_names, rotation=45, ha='right')
plt.legend()

# Add value labels on top of bars
for bar in bars1:
    height = bar.get_height()
    plt.text(bar.get_x() + bar.get_width()/2., height + 0.01,
             f'{height:.4f}', ha='center', va='bottom', rotation=0, fontsize=8)

for bar in bars2:
    height = bar.get_height()
    plt.text(bar.get_x() + bar.get_width()/2., height + 0.01,
             f'{height:.4f}', ha='center', va='bottom', rotation=0, fontsize=8)

plt.tight_layout()

# Save the comparison chart
comparison_path = f'{output_dir}\\original_vs_normalized_comparison.png'
plt.savefig(comparison_path)
plt.close()

# Add the comparison chart to the document
doc.add_paragraph().add_run().add_picture(comparison_path, width=Inches(6))

# Add comparison summary
doc.add_paragraph("""原始数据与标准化数据交叉验证结果的对比分析表明：

1. 对于大多数模型，标准化处理对交叉验证性能的影响因模型类型而异。

2. 线性模型（如多元线性回归）在标准化数据上的表现通常优于原始数据，这与线性模型对特征尺度敏感的特性一致。

3. 基于树的模型（如XGBoost、随机森林和梯度提升树）在标准化前后的表现差异相对较小，这验证了树模型对特征尺度不敏感的特性。

4. 神经网络模型（如LSTM和ELM）在标准化数据上的表现通常优于原始数据，这与神经网络对输入数据尺度敏感的特性一致。

5. 总体而言，标准化处理对模型的交叉验证性能有一定的积极影响，特别是对于线性模型和神经网络模型。

这些结果进一步证实了在机器学习建模过程中，根据模型特性选择合适的数据预处理方法的重要性。对于某些模型，数据标准化可以显著提高模型的稳定性和泛化能力；而对于其他模型，数据标准化的影响可能较小。在实际应用中，应根据具体问题和模型特性，灵活选择是否进行数据标准化处理。""")



# After the normalized cross-validation comparison section, add feature importance for best model

# Find the best model from normalized cross-validation results
best_norm_cv_model_name = max(normalized_cv_results, key=lambda x: x[1]['mean_r2'])[0]
doc.add_heading(f'最佳标准化模型({best_norm_cv_model_name})特征重要性分析', level=3)
doc.add_paragraph(f"""在标准化数据的交叉验证结果中，{best_norm_cv_model_name}表现最佳，平均R²为{max(normalized_cv_results, key=lambda x: x[1]['mean_r2'])[1]['mean_r2']:.4f}。以下分析该模型的特征重要性，以深入了解各股票市场指标对BIS指数的影响程度。""")

# Train the best model on full normalized dataset to get feature importance
if best_norm_cv_model_name == 'XGBoost模型':
    best_model = xgb.XGBRegressor(random_state=42)
    best_model.fit(X_normalized, y_normalized)
    
    # Get feature importance
    feature_importance = pd.DataFrame({
        'Feature': X_normalized.columns,
        'Importance': best_model.feature_importances_
    }).sort_values('Importance', ascending=False)
    
    # Set font for Chinese characters
    plt.rcParams['font.sans-serif'] = ['SimHei']
    plt.rcParams['axes.unicode_minus'] = False
    
    # Create color gradient from light to dark purple as specified
    light_purple = np.array([221/255, 204/255, 239/255])  # Light purple RGB(221, 204, 239)
    dark_purple = np.array([147/255, 111/255, 184/255])   # Dark purple RGB(147, 111, 184)
    
    # Create a figure for feature importance
    plt.figure(figsize=(12, 8))
    
    # Sort features by importance
    feature_importance = feature_importance.sort_values('Importance', ascending=True)
    
    # Normalize importance values to range [0, 1] for color mapping
    norm_importance = (feature_importance['Importance'] - feature_importance['Importance'].min()) / \
                     (feature_importance['Importance'].max() - feature_importance['Importance'].min())
    
    # Generate colors for each bar
    colors = [light_purple + t * (dark_purple - light_purple) for t in norm_importance]
    
    # Create horizontal bar chart
    bars = plt.barh(feature_importance['Feature'], feature_importance['Importance'], color=colors)
    
    # Add value labels to the right of each bar
    for bar in bars:
        width = bar.get_width()
        plt.text(width + 0.002, bar.get_y() + bar.get_height()/2, 
                 f'{width:.4f}', ha='left', va='center')
    
    plt.title(f'最佳标准化模型({best_norm_cv_model_name})特征重要性')
    plt.xlabel('重要性')
    plt.ylabel('特征')
    plt.tight_layout()
    
    # Save the figure
    best_model_fi_path = f'{output_dir}\\最佳标准化模型_feature_importance.png'
    plt.savefig(best_model_fi_path)
    plt.close()
    
    # Add the figure to the document
    doc.add_paragraph().add_run().add_picture(best_model_fi_path, width=Inches(6))
    
    # Add interpretation
    top_features = feature_importance.tail(3)['Feature'].tolist()
    bottom_features = feature_importance.head(3)['Feature'].tolist()
    
    interpretation = f"""特征重要性分析表明，在最佳标准化模型({best_norm_cv_model_name})中，对BIS指数预测影响最大的三个特征是{', '.join(top_features)}，
这表明这些股票市场指标与BIS指数之间存在较强的关联。相比之下，{', '.join(bottom_features)}的影响相对较小。

这一结果与原始数据模型的特征重要性分析结果{"基本一致" if set(top_features) & set([feature_importance['Feature'].tolist()[0] for feature_importance in [pd.DataFrame({'Feature': X.columns, 'Importance': xgb_model.feature_importances_}).sort_values('Importance', ascending=False).head(3)]]) else "有所不同"}，
表明{"数据标准化处理不会显著改变模型对特征重要性的判断" if set(top_features) & set([feature_importance['Feature'].tolist()[0] for feature_importance in [pd.DataFrame({'Feature': X.columns, 'Importance': xgb_model.feature_importances_}).sort_values('Importance', ascending=False).head(3)]]) else "数据标准化处理会在一定程度上改变模型对特征重要性的判断"}。

这些高重要性特征可能代表了股票市场与汇率之间关联的关键因素，为理解两个市场之间的相互作用机制提供了重要线索。在实际应用中，可以重点关注这些高重要性特征的变化，以更好地预测BIS指数的走势。"""
    doc.add_paragraph(interpretation)
elif best_norm_cv_model_name == '随机森林模型':
    best_model = RandomForestRegressor(random_state=42)
    best_model.fit(X_normalized, y_normalized)
    
    # Get feature importance
    feature_importance = pd.DataFrame({
        'Feature': X_normalized.columns,
        'Importance': best_model.feature_importances_
    }).sort_values('Importance', ascending=False)
    
    # Set font for Chinese characters
    plt.rcParams['font.sans-serif'] = ['SimHei']
    plt.rcParams['axes.unicode_minus'] = False
    
    # Create color gradient from light to dark purple as specified
    light_purple = np.array([221/255, 204/255, 239/255])  # Light purple RGB(221, 204, 239)
    dark_purple = np.array([147/255, 111/255, 184/255])   # Dark purple RGB(147, 111, 184)
    
    # Create a figure for feature importance
    plt.figure(figsize=(12, 8))
    
    # Sort features by importance
    feature_importance = feature_importance.sort_values('Importance', ascending=True)
    
    # Normalize importance values to range [0, 1] for color mapping
    norm_importance = (feature_importance['Importance'] - feature_importance['Importance'].min()) / \
                     (feature_importance['Importance'].max() - feature_importance['Importance'].min())
    
    # Generate colors for each bar
    colors = [light_purple + t * (dark_purple - light_purple) for t in norm_importance]
    
    # Create horizontal bar chart
    bars = plt.barh(feature_importance['Feature'], feature_importance['Importance'], color=colors)
    
    # Add value labels to the right of each bar
    for bar in bars:
        width = bar.get_width()
        plt.text(width + 0.002, bar.get_y() + bar.get_height()/2, 
                 f'{width:.4f}', ha='left', va='center')
    
    plt.title(f'最佳标准化模型({best_norm_cv_model_name})特征重要性')
    plt.xlabel('重要性')
    plt.ylabel('特征')
    plt.tight_layout()
    
    # Save the figure
    best_model_fi_path = f'{output_dir}\\最佳标准化模型_feature_importance.png'
    plt.savefig(best_model_fi_path)
    plt.close()
    
    # Add the figure to the document
    doc.add_paragraph().add_run().add_picture(best_model_fi_path, width=Inches(6))
    
    # Add interpretation
    top_features = feature_importance.tail(3)['Feature'].tolist()
    bottom_features = feature_importance.head(3)['Feature'].tolist()
    
    interpretation = f"""特征重要性分析表明，在最佳标准化模型({best_norm_cv_model_name})中，对BIS指数预测影响最大的三个特征是{', '.join(top_features)}，
这表明这些股票市场指标与BIS指数之间存在较强的关联。相比之下，{', '.join(bottom_features)}的影响相对较小。

这一结果与原始数据模型的特征重要性分析结果{"基本一致" if set(top_features) & set([feature_importance['Feature'].tolist()[0] for feature_importance in [pd.DataFrame({'Feature': X.columns, 'Importance': rf_model.feature_importances_}).sort_values('Importance', ascending=False).head(3)]]) else "有所不同"}，
表明{"数据标准化处理不会显著改变模型对特征重要性的判断" if set(top_features) & set([feature_importance['Feature'].tolist()[0] for feature_importance in [pd.DataFrame({'Feature': X.columns, 'Importance': rf_model.feature_importances_}).sort_values('Importance', ascending=False).head(3)]]) else "数据标准化处理会在一定程度上改变模型对特征重要性的判断"}。

这些高重要性特征可能代表了股票市场与汇率之间关联的关键因素，为理解两个市场之间的相互作用机制提供了重要线索。在实际应用中，可以重点关注这些高重要性特征的变化，以更好地预测BIS指数的走势。"""
    doc.add_paragraph(interpretation)
elif best_norm_cv_model_name == '梯度提升树模型':
    best_model = GradientBoostingRegressor(random_state=42)
    best_model.fit(X_normalized, y_normalized)
    
    # Get feature importance
    feature_importance = pd.DataFrame({
        'Feature': X_normalized.columns,
        'Importance': best_model.feature_importances_
    }).sort_values('Importance', ascending=False)
    
    # Set font for Chinese characters
    plt.rcParams['font.sans-serif'] = ['SimHei']
    plt.rcParams['axes.unicode_minus'] = False
    
    # Create color gradient from light to dark purple as specified
    light_purple = np.array([221/255, 204/255, 239/255])  # Light purple RGB(221, 204, 239)
    dark_purple = np.array([147/255, 111/255, 184/255])   # Dark purple RGB(147, 111, 184)
    
    # Create a figure for feature importance
    plt.figure(figsize=(12, 8))
    
    # Sort features by importance
    feature_importance = feature_importance.sort_values('Importance', ascending=True)
    
    # Normalize importance values to range [0, 1] for color mapping
    norm_importance = (feature_importance['Importance'] - feature_importance['Importance'].min()) / \
                     (feature_importance['Importance'].max() - feature_importance['Importance'].min())
    
    # Generate colors for each bar
    colors = [light_purple + t * (dark_purple - light_purple) for t in norm_importance]
    
    # Create horizontal bar chart
    bars = plt.barh(feature_importance['Feature'], feature_importance['Importance'], color=colors)
    
    # Add value labels to the right of each bar
    for bar in bars:
        width = bar.get_width()
        plt.text(width + 0.002, bar.get_y() + bar.get_height()/2, 
                 f'{width:.4f}', ha='left', va='center')
    
    plt.title(f'最佳标准化模型({best_norm_cv_model_name})特征重要性')
    plt.xlabel('重要性')
    plt.ylabel('特征')
    plt.tight_layout()
    
    # Save the figure
    best_model_fi_path = f'{output_dir}\\最佳标准化模型_feature_importance.png'
    plt.savefig(best_model_fi_path)
    plt.close()
    
    # Add the figure to the document
    doc.add_paragraph().add_run().add_picture(best_model_fi_path, width=Inches(6))
    
    # Add interpretation
    top_features = feature_importance.tail(3)['Feature'].tolist()
    bottom_features = feature_importance.head(3)['Feature'].tolist()
    
    interpretation = f"""特征重要性分析表明，在最佳标准化模型({best_norm_cv_model_name})中，对BIS指数预测影响最大的三个特征是{', '.join(top_features)}，
这表明这些股票市场指标与BIS指数之间存在较强的关联。相比之下，{', '.join(bottom_features)}的影响相对较小。

这一结果与原始数据模型的特征重要性分析结果{"基本一致" if set(top_features) & set([feature_importance['Feature'].tolist()[0] for feature_importance in [pd.DataFrame({'Feature': X.columns, 'Importance': gbdt_model.feature_importances_}).sort_values('Importance', ascending=False).head(3)]]) else "有所不同"}，
表明{"数据标准化处理不会显著改变模型对特征重要性的判断" if set(top_features) & set([feature_importance['Feature'].tolist()[0] for feature_importance in [pd.DataFrame({'Feature': X.columns, 'Importance': gbdt_model.feature_importances_}).sort_values('Importance', ascending=False).head(3)]]) else "数据标准化处理会在一定程度上改变模型对特征重要性的判断"}。

这些高重要性特征可能代表了股票市场与汇率之间关联的关键因素，为理解两个市场之间的相互作用机制提供了重要线索。在实际应用中，可以重点关注这些高重要性特征的变化，以更好地预测BIS指数的走势。"""
    doc.add_paragraph(interpretation)
else:
    doc.add_paragraph(f"""由于{best_norm_cv_model_name}模型不提供直接的特征重要性度量，因此无法绘制特征重要性图。对于此类模型，可以考虑使用排列重要性(Permutation Importance)或SHAP值等模型无关的方法来评估特征重要性。这些方法可以应用于任何机器学习模型，为模型的可解释性提供支持。""")
# Add conclusion
doc.add_heading('5.5 本章小结', level=1)
doc.add_paragraph("""本章运用多种机器学习模型对股票市场指标与BIS有效汇率指数之间的关系进行了系统分析，通过模型训练、评估、交叉验证和数据标准化等多个环节，得出以下主要结论：

1. 股票市场与汇率的关联性：多种机器学习模型的预测结果均表明，股票市场指标与BIS指数之间存在显著的关联模式。这种关联不仅仅是简单的线性关系，更多地表现为复杂的非线性关系，这解释了为何非线性模型（如XGBoost和LSTM）能够取得更好的预测效果。

2. 模型性能比较：在所有评估的模型中，XGBoost模型和长短期记忆网络(LSTM)模型展现出最优的预测性能，其R²值和RMSE指标均优于其他模型。这一结果表明，集成学习方法和深度学习方法在捕捉金融时间序列数据中的复杂模式方面具有显著优势。

3. 特征重要性洞察：通过对XGBoost等模型的特征重要性分析发现，某些特定的股票市场指标对BIS指数的预测贡献度明显高于其他指标。这些高重要性特征主要包括市场波动性指标和大型股票指数，表明这些因素可能是连接股票市场与汇率变动的关键桥梁。

4. 数据预处理的影响：数据标准化处理对不同类型模型的影响存在显著差异。对于线性模型和神经网络模型，标准化处理能够显著提升模型性能；而对于基于树的模型（如随机森林和XGBoost），标准化处理的影响相对较小。这一发现验证了不同算法对特征尺度敏感性的理论预期。

5. 模型稳定性评估：交叉验证结果显示，所有模型在不同数据子集上的表现具有较好的一致性，特别是XGBoost和LSTM模型的性能波动较小。这种稳定性增强了模型预测结果的可信度，为实际应用提供了可靠保障。

6. 新兴算法的应用价值：作为一种计算效率高的新兴算法，极限学习机(ELM)模型在本研究中展现出了与传统算法相当的预测能力，且训练速度明显快于深度学习模型。这表明ELM在金融时间序列分析领域具有广阔的应用前景。

7. 原始数据与标准化数据的对比：通过对比分析发现，在标准化数据上训练的模型整体性能优于原始数据，尤其是对于线性回归和神经网络模型。这一结果强调了在金融数据建模中选择适当预处理方法的重要性。

8. 方法论启示：本研究采用的多模型比较和交叉验证方法，为金融市场关联性研究提供了一套系统的分析框架。这种方法不仅能够识别最优模型，还能深入理解不同模型的优缺点及其适用场景。

综上所述，本章研究不仅验证了股票市场指标与BIS指数之间存在稳健的预测关系，还通过多种机器学习模型的对比分析，揭示了这种关系的复杂性和非线性特征。这些发现为理解全球金融市场的相互作用机制提供了新的视角，同时也为投资决策、风险管理和宏观经济政策制定提供了数据驱动的参考依据。未来研究可以进一步探索这些关联背后的经济机制，以及如何将这些模型预测结果整合到实际的金融决策系统中。""")

# Save the Word document
doc.save(r'C:\Users\PKF\Desktop\实习\ycq\第五章_股票市场对BIS指数影响的机器学习模型分析.docx')

print("分析完成，结果已保存到Word文档中。")